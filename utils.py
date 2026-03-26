import os
import time
import zipfile
from xml.etree import ElementTree

import PyPDF2
from werkzeug.utils import secure_filename

from config import Config


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in Config.ALLOWED_EXTENSIONS


def get_file_extension(file_path):
    return os.path.splitext(file_path)[1].lower().lstrip('.')


def extract_text_from_pdf(file_path):
    try:
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text() or '')
        return '\n'.join(part.strip() for part in text if part and part.strip()).strip()
    except Exception as exc:
        raise Exception(f'PDF 解析失败: {exc}')


def extract_text_from_docx(file_path):
    try:
        from docx import Document
    except ImportError as exc:
        raise Exception('缺少 python-docx 依赖，请先执行 pip install -r requirements.txt') from exc

    try:
        document = Document(file_path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

        header_footer_text = []
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                text = paragraph.text.strip()
                if text:
                    header_footer_text.append(text)
            for paragraph in section.footer.paragraphs:
                text = paragraph.text.strip()
                if text:
                    header_footer_text.append(text)

        tables = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    tables.append(' | '.join(cells))

        combined = '\n'.join(paragraphs + header_footer_text + tables).strip()
        if combined:
            return combined

        xml_text = extract_text_from_docx_xml(file_path)
        if xml_text:
            return xml_text

        raise Exception('DOCX 文件中未提取到可识别文本')
    except Exception as exc:
        raise Exception(f'DOCX 解析失败: {exc}')


def extract_text_from_docx_xml(file_path):
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    parts = []

    with zipfile.ZipFile(file_path) as archive:
        for name in archive.namelist():
            if not name.startswith('word/') or not name.endswith('.xml'):
                continue
            try:
                root = ElementTree.fromstring(archive.read(name))
            except ElementTree.ParseError:
                continue

            texts = []
            for node in root.findall('.//w:t', namespace):
                text = ''.join(node.itertext()).strip()
                if text:
                    texts.append(text)

            if texts:
                parts.append('\n'.join(texts))

    deduplicated = []
    for chunk in parts:
        normalized = chunk.strip()
        if normalized and normalized not in deduplicated:
            deduplicated.append(normalized)

    return '\n'.join(deduplicated).strip()


def extract_text_from_file(file_path):
    extension = get_file_extension(file_path)
    if extension == 'pdf':
        return extract_text_from_pdf(file_path)
    if extension == 'docx':
        return extract_text_from_docx(file_path)
    raise Exception(f'暂不支持解析 .{extension} 文件')


def save_uploaded_file(file, upload_folder):
    if not os.path.exists(upload_folder):
        os.makedirs(upload_folder)

    filename = secure_filename(file.filename)
    timestamp = int(time.time())
    filename = f'{timestamp}_{filename}'
    file_path = os.path.join(upload_folder, filename)
    file.save(file_path)
    return filename, file_path
