import os
from datetime import datetime


def ensure_directory(path):
    os.makedirs(path, exist_ok=True)
    return path


def export_optimization(version, export_format, export_folder):
    export_format = (export_format or 'txt').lower()
    if export_format not in {'txt', 'md', 'pdf', 'docx'}:
        raise ValueError('暂仅支持 txt、md、pdf 和 docx 导出')

    ensure_directory(export_folder)
    timestamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
    safe_title = ''.join(char if char.isalnum() else '_' for char in version.title)[:60].strip('_') or 'optimized_resume'
    filename = f'{safe_title}_{timestamp}.{export_format}'
    file_path = os.path.join(export_folder, filename)

    if export_format in {'txt', 'md'}:
        body = build_export_body(version, export_format)
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(body)
        return file_path

    if export_format == 'docx':
        export_docx(version, file_path)
        return file_path

    export_pdf(version, file_path)
    return file_path


def build_export_body(version, export_format):
    if export_format == 'md':
        lines = [
            f'# {version.title}',
            '',
            f'- 版本类型：{version.version_type}',
            f'- 目标岗位：{version.target_job_title or "未指定"}',
            '',
            '## 优化摘要',
            version.summary or '暂无摘要',
            '',
            '## 优化亮点',
        ]
        highlights = version.highlights or ['暂无亮点']
        lines.extend([f'- {item}' for item in highlights])
        lines.extend(['', '## 优化后内容', version.content])
        return '\n'.join(lines)

    lines = [
        version.title,
        '=' * len(version.title),
        f'版本类型：{version.version_type}',
        f'目标岗位：{version.target_job_title or "未指定"}',
        '',
        '优化摘要：',
        version.summary or '暂无摘要',
        '',
        '优化亮点：',
    ]
    highlights = version.highlights or ['暂无亮点']
    lines.extend([f'{index + 1}. {item}' for index, item in enumerate(highlights)])
    lines.extend(['', '优化后内容：', version.content])
    return '\n'.join(lines)


def export_docx(version, file_path):
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError('缺少 python-docx 依赖，请先执行 pip install -r requirements.txt') from exc

    document = Document()
    document.add_heading(version.title, level=1)
    document.add_paragraph(f'版本类型：{version.version_type}')
    document.add_paragraph(f'目标岗位：{version.target_job_title or "未指定"}')
    document.add_heading('优化摘要', level=2)
    document.add_paragraph(version.summary or '暂无摘要')
    document.add_heading('优化亮点', level=2)
    for item in version.highlights or ['暂无亮点']:
        document.add_paragraph(item, style='List Bullet')
    document.add_heading('优化后内容', level=2)
    for paragraph in (version.content or '').splitlines():
        document.add_paragraph(paragraph or '')
    document.save(file_path)


def export_pdf(version, file_path):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise ValueError('缺少 reportlab 依赖，请先执行 pip install -r requirements.txt') from exc

    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    pdf = canvas.Canvas(file_path, pagesize=A4)
    width, height = A4
    x = 48
    y = height - 48
    line_height = 18

    pdf.setFont('STSong-Light', 18)
    pdf.drawString(x, y, version.title)
    y -= line_height * 2

    lines = [
        f'版本类型：{version.version_type}',
        f'目标岗位：{version.target_job_title or "未指定"}',
        '',
        '优化摘要：',
        version.summary or '暂无摘要',
        '',
        '优化亮点：',
    ]
    lines.extend([f'{index + 1}. {item}' for index, item in enumerate(version.highlights or ['暂无亮点'])])
    lines.extend(['', '优化后内容：'])
    lines.extend((version.content or '').splitlines())

    pdf.setFont('STSong-Light', 11)
    for raw_line in lines:
        wrapped_lines = wrap_pdf_line(raw_line, 42)
        for line in wrapped_lines:
            if y < 60:
                pdf.showPage()
                pdf.setFont('STSong-Light', 11)
                y = height - 48
            pdf.drawString(x, y, line)
            y -= line_height

    pdf.save()


def wrap_pdf_line(text, width):
    if not text:
        return ['']
    lines = []
    current = ''
    for char in text:
        current += char
        if len(current) >= width:
            lines.append(current)
            current = ''
    if current:
        lines.append(current)
    return lines or ['']
